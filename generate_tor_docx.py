"""Generate TOR_SIAM_EV_PLATFORM.docx from structured data using python-docx."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "TH Sarabun New"
style.font.size = Pt(14)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "TH Sarabun New"
    h.font.bold = True
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)
    h.font.size = Pt({1: 22, 2: 18, 3: 15}[level])
    h.paragraph_format.space_before = Pt({1: 24, 2: 18, 3: 12}[level])
    h.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "TH Sarabun New"
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "TH Sarabun New"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A5676")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "TH Sarabun New"
            run.font.size = Pt(12)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF5FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table


# ═══════════════════════════════════════════════════════════════
#  COVER / TITLE
# ═══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

add_para("ข้อกำหนดขอบเขตงาน", bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Terms of Reference (TOR)", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_para("ระบบบริหารจัดการพลังงานและโลจิสติกส์อัจฉริยะครบวงจร", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("SIAM EV — Unified Energy & Logistics Management Platform", bold=False, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_para("เวอร์ชัน TOR: 1.0", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("วันที่จัดทำ: เมษายน 2569", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════
doc.add_heading("สารบัญ", level=1)
toc_items = [
    "1. บทนำและวัตถุประสงค์",
    "2. สถาปัตยกรรมระบบ",
    "3. โมดูล 1 — EV Charging Management",
    "4. โมดูล 2 — Fleet Optimization",
    "5. โมดูล 3 — Facilities Energy Analytics",
    "6. โมดูล 4 — Cross-Domain Intelligence",
    "7. โมดูล 5 — Platform Foundation",
    "8. Mobile Application",
    "9. AI/ML Capabilities",
    "10. Non-Functional Requirements",
    "11. แผนการพัฒนา (Roadmap)",
    "12. เกณฑ์การตัดสิน",
    "13. ตารางสรุปฟังก์ชันทั้งหมด",
    "14. คำจำกัดความและคำย่อ",
]
for item in toc_items:
    add_para(item, size=14, space_after=2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading("1. บทนำและวัตถุประสงค์", level=1)

doc.add_heading("1.1 ภาพรวม", level=2)
add_para("ระบบ SIAM EV Platform เป็นแพลตฟอร์มบริหารจัดการครบวงจร (Unified Operations Platform) ที่ผนวก 3 โดเมนหลักเข้าด้วยกัน:")
add_table(
    ["โดเมน", "ขอบเขต"],
    [
        ["EV Charging Management", "บริหารสถานีชาร์จ EV, รอบการชาร์จ, Smart Charging, V2G"],
        ["Fleet Optimization", "จัดเส้นทาง, ติดตามยานพาหนะ, วางแผนการชาร์จฟลีท EV"],
        ["Facilities Energy Analytics", "วิเคราะห์การใช้พลังงานอาคาร, Demand Response, พลังงานหมุนเวียน"],
    ],
    col_widths=[5, 12],
)

doc.add_heading("1.2 วัตถุประสงค์", level=2)
objectives = [
    "รวมศูนย์การบริหารจัดการสถานีชาร์จ, กองยานพาหนะ, และพลังงานอาคาร บนแพลตฟอร์มเดียว",
    "ลดต้นทุนพลังงานรวมขององค์กรผ่าน Cross-Domain Optimization",
    "เพิ่มประสิทธิภาพการใช้ทรัพยากรด้วย AI/ML Predictive Analytics",
    "รองรับมาตรฐานสากล OCPP 2.0.1, OCPI 2.2, ISO 15118",
]
for i, obj in enumerate(objectives, 1):
    add_para(f"{i}. {obj}")

doc.add_heading("1.3 ขอบเขตการให้บริการ", level=2)
add_table(
    ["รายการ", "ขอบเขต"],
    [
        ["พื้นที่นำร่อง", "กรุงเทพมหานครและปริมณฑล"],
        ["สถานีชาร์จ", "รองรับ 10–500 สถานี"],
        ["กองยาน", "รองรับ 6–200 คัน (EV และ ICE)"],
        ["อาคาร/สิ่งอำนวยความสะดวก", "รองรับ 5–50 แห่ง"],
    ],
    col_widths=[5, 12],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading("2. สถาปัตยกรรมระบบ (System Architecture)", level=1)

doc.add_heading("2.1 ภาพรวมสถาปัตยกรรม", level=2)
add_para("ระบบออกแบบเป็น 3 ชั้น (Three-tier Architecture):")
add_para("• Presentation Layer — Web Dashboard (Next.js), Mobile App (React Native), Kiosk Mode, Public API")
add_para("• Application Layer — API Gateway, Microservices (EV / Fleet / Facilities / Cross-Domain / Auth / Notification / Reporting / AI-ML)")
add_para("• Data & Integration Layer — PostgreSQL + TimescaleDB, Redis, MQTT Broker, Node-RED (OCPP Gateway)")

doc.add_heading("2.2 เทคโนโลยีหลัก", level=2)
add_table(
    ["ชั้น", "เทคโนโลยี"],
    [
        ["Frontend", "Next.js 16, React 19, TypeScript, Tailwind CSS v4, Leaflet/MapLibre"],
        ["Backend", "Node.js (NestJS), Python (FastAPI สำหรับ ML pipeline)"],
        ["Real-time", "WebSocket, MQTT (Mosquitto), Server-Sent Events"],
        ["Database", "PostgreSQL 16 + TimescaleDB (time-series), Redis (cache & pub/sub)"],
        ["OCPP Gateway", "Node-RED + node-red-contrib-ocpp"],
        ["ML/AI", "Python, scikit-learn, Prophet, TensorFlow Lite"],
        ["Deployment", "Docker Compose, Kubernetes (production), CapRover (staging)"],
        ["CI/CD", "GitHub Actions"],
    ],
    col_widths=[4, 13],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  3. EV CHARGING MANAGEMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading("3. โมดูล 1 — EV Charging Management (CSMS)", level=1)

doc.add_heading("3.1 ฟังก์ชันที่มีในระบบปัจจุบัน (Baseline)", level=2)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด"],
    [
        ["EV-B01", "แผนที่สถานีชาร์จ Real-time", "แสดงตำแหน่งสถานีบน Leaflet Map พร้อมสถานะ (Available/Charging/Faulted)"],
        ["EV-B02", "Monitoring Dashboard", "แสดง Net Energy (kWh), Max Capacity (kW), Live Output (kW) แบบ Real-time"],
        ["EV-B03", "Station Detail Panel", "แสดงรายละเอียดสถานี: กำลังไฟ, สถานะซ็อกเก็ต, มิเตอร์สะสม"],
        ["EV-B04", "Schedule Management", "ตั้งเวลาเปิด-ปิดสถานีชาร์จ (startTime/endTime/enabled)"],
        ["EV-B05", "OCPP Proxy", "Reverse proxy เชื่อมต่อ Node-RED OCPP backend"],
        ["EV-B06", "Nodes Table View", "ตารางสถานีชาร์จทั้งหมด พร้อมการค้นหาและเรียงลำดับ"],
    ],
    col_widths=[2, 4, 11],
)

doc.add_heading("3.2 ฟังก์ชันใหม่ (Enhanced)", level=2)

doc.add_heading("3.2.1 OCPP 2.0.1 Full Compliance", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["EV-E01", "OCPP 2.0.1 Core Profile", "รองรับ Boot Notification, Heartbeat, Status Notification, Authorize, Start/Stop Transaction, Meter Values", "P0"],
        ["EV-E02", "Remote Start/Stop", "สั่ง Start/Stop Transaction จาก Dashboard ผ่าน OCPP", "P0"],
        ["EV-E03", "Firmware Management", "OTA firmware update: upload, schedule, rollback", "P1"],
        ["EV-E04", "Diagnostics Collection", "ดึง diagnostics logs จากสถานีชาร์จ", "P1"],
        ["EV-E05", "Configuration Management", "อ่าน/เขียน Configuration Key ของสถานีชาร์จ", "P1"],
        ["EV-E06", "Reservation Support", "จองสถานีชาร์จล่วงหน้า (ReserveNow, CancelReservation)", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("3.2.2 Smart Charging", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["EV-S01", "Dynamic Load Balancing", "กระจายโหลดอัตโนมัติระหว่างสถานีชาร์จเมื่อกำลังไฟรวมเกิน Site Power Limit", "P0"],
        ["EV-S02", "Charging Profile Mgmt", "สร้างและจัดการ Charging Profile ตาม OCPP Smart Charging", "P0"],
        ["EV-S03", "TOU Optimization", "จัดตารางชาร์จตามอัตราค่าไฟ TOU (On-Peak/Off-Peak/Holiday)", "P1"],
        ["EV-S04", "Priority Queue", "จัดลำดับการชาร์จ: รถฟลีทที่ต้องออกเร็ว > รถทั่วไป > รถชาร์จเต็ม", "P1"],
        ["EV-S05", "V2G Readiness", "รองรับ bidirectional charging: V2G capability, discharge schedule, คำนวณรายได้", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("3.2.3 Session & Billing", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["EV-SB01", "Session Lifecycle", "บันทึกครบวงจร: เริ่มชาร์จ → Meter Values → หยุดชาร์จ → สรุปค่าใช้จ่าย", "P0"],
        ["EV-SB02", "Transaction History", "ประวัติธุรกรรมย้อนหลัง 12 เดือน พร้อมกรอง", "P0"],
        ["EV-SB03", "Dynamic Pricing", "คำนวณค่าชาร์จ: ราคาตาม kWh, TOU, demand surcharge, สมาชิก discount", "P1"],
        ["EV-SB04", "CDR Generation (OCPI)", "สร้าง Charge Detail Record ตามมาตรฐาน OCPI 2.2", "P2"],
        ["EV-SB05", "Payment Integration", "QR PromptPay, บัตรเครดิต, RFID card, mobile wallet", "P1"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("3.2.4 Predictive Maintenance", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["EV-PM01", "Health Score", "คะแนนสุขภาพ 0–100 จาก: uptime, error frequency, power variance, temperature", "P1"],
        ["EV-PM02", "Anomaly Detection", "ตรวจจับ: แรงดันผิดปกติ, current leakage, temperature spike", "P1"],
        ["EV-PM03", "Failure Prediction", "คาดการณ์สถานีจะเสียภายในกี่วัน (MTBF prediction)", "P2"],
        ["EV-PM04", "Maintenance Scheduler", "สร้างตาราง preventive maintenance อัตโนมัติ", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  4. FLEET OPTIMIZATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading("4. โมดูล 2 — Fleet Optimization", level=1)

doc.add_heading("4.1 ฟังก์ชันที่มีในระบบปัจจุบัน (Baseline)", level=2)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด"],
    [
        ["FL-B01", "แผนที่ติดตามยานพาหนะ", "แสดงตำแหน่ง, ทิศทาง, ความเร็ว Real-time บนแผนที่"],
        ["FL-B02", "สถานะยานพาหนะ", "En Route / Idle / Loading / Returning / Maintenance"],
        ["FL-B03", "ข้อมูลเส้นทาง", "origin, destination, waypoints, ETA, ระยะทางเหลือ"],
        ["FL-B04", "ข้อมูลสินค้า", "จำนวนพัสดุ: loaded, capacity, available, load level (0–5)"],
        ["FL-B05", "Fleet Table View", "ตารางยานพาหนะทั้งหมด พร้อมข้อมูลคนขับ"],
    ],
    col_widths=[2, 4, 11],
)

doc.add_heading("4.2 ฟังก์ชันใหม่ (Enhanced)", level=2)

doc.add_heading("4.2.1 Route Optimization", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FL-R01", "Multi-Stop Route Optimization", "คำนวณเส้นทางประหยัดสุดสำหรับหลายจุด (OR-Tools / Google Directions API)", "P0"],
        ["FL-R02", "Real-time Re-routing", "ปรับเส้นทางอัตโนมัติเมื่อเจอจราจร/อุบัติเหตุ", "P1"],
        ["FL-R03", "Geofencing & Alerts", "กำหนดเขตพื้นที่: แจ้งเตือนเข้า/ออก, คำนวณ dwell time", "P1"],
        ["FL-R04", "Historical Route Analytics", "เวลาจริง vs ETA, เส้นทางที่ใช้บ่อย, คอขวดจราจร", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("4.2.2 EV Fleet Charging Orchestration", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FL-EC01", "Battery State Tracking", "ติดตาม SoC, SoH, estimated range สำหรับรถ EV", "P0"],
        ["FL-EC02", "Charge Planning", "วางแผนชาร์จ: SoC ขั้นต่ำ, ตาราง TOU, ความพร้อมสถานี", "P0"],
        ["FL-EC03", "Range Anxiety Prevention", "แจ้งเตือนเมื่อ SoC ไม่พอถึงจุดหมาย + แนะนำสถานีระหว่างทาง", "P1"],
        ["FL-EC04", "Depot Charging Scheduler", "ตารางชาร์จกลางคืน: minimize TOU + ทุกคันพร้อมก่อนออกงาน", "P1"],
        ["FL-EC05", "Battery Degradation", "ติดตาม SoH trend, คาดการณ์อายุแบต, แนะนำช่วงเปลี่ยน", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("4.2.3 Dispatch & Load Optimization", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FL-D01", "Smart Dispatch", "กระจายงานตาม: ตำแหน่งใกล้สุด, capacity, SoC, ชั่วโมงคนขับ", "P0"],
        ["FL-D02", "Load Consolidation", "รวม shipments ทิศทางเดียวกันเข้ารถคันเดียว", "P1"],
        ["FL-D03", "Driver Performance Score", "eco-driving score, on-time rate, energy efficiency", "P1"],
        ["FL-D04", "Driver HOS Compliance", "ติดตามชั่วโมงทำงานคนขับ ป้องกันทำงานเกินเวลา", "P2"],
        ["FL-D05", "Proof of Delivery", "ภาพถ่าย, ลายเซ็นดิจิทัล, timestamp, GPS", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("4.2.4 Fleet Analytics", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FL-A01", "Fleet Utilization Dashboard", "idle time, active time, utilization %, km/วัน", "P0"],
        ["FL-A02", "TCO Analysis", "ค่าพลังงาน, บำรุงรักษา, ค่าเสื่อม, ประกัน ต่อคัน ต่อ km", "P1"],
        ["FL-A03", "EV vs ICE Comparison", "เปรียบเทียบ TCO เพื่อวางแผน fleet transition", "P2"],
        ["FL-A04", "Carbon Footprint Tracker", "คำนวณ CO2 ที่ลดได้จาก EV fleet เทียบกับ ICE baseline", "P1"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  5. FACILITIES ENERGY ANALYTICS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("5. โมดูล 3 — Facilities Energy Analytics", level=1)

doc.add_heading("5.1 ฟังก์ชันที่มีในระบบปัจจุบัน (Baseline)", level=2)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด"],
    [
        ["FA-B01", "แผนที่สิ่งอำนวยความสะดวก", "แสดงตำแหน่ง Warehouse, Office, Factory, Distribution Center"],
        ["FA-B02", "Real-time Power Monitoring", "กำลังไฟ (kW), พลังงานรายวัน/รายเดือน (kWh), Peak Demand, PUE"],
        ["FA-B03", "Zone-level Breakdown", "การใช้พลังงานแยกตามโซน: ชั้น, สำนักงาน, คลังสินค้า, ห้องเย็น"],
        ["FA-B04", "สถานะอาคาร", "Normal / High Usage / Critical / Offline"],
        ["FA-B05", "Occupancy Tracking", "% การใช้พื้นที่ของแต่ละโซน"],
    ],
    col_widths=[2, 4, 11],
)

doc.add_heading("5.2 ฟังก์ชันใหม่ (Enhanced)", level=2)

doc.add_heading("5.2.1 Advanced Energy Monitoring", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FA-M01", "Sub-metering Integration", "เชื่อมต่อ smart meter (Modbus/BACnet/MQTT) แยกตาม HVAC, Lighting, Equipment", "P0"],
        ["FA-M02", "Power Quality Monitoring", "voltage, current, power factor, harmonic distortion", "P1"],
        ["FA-M03", "Energy Intensity Metrics", "kWh/m², kWh/unit produced, kWh/employee สำหรับ benchmarking", "P1"],
        ["FA-M04", "Cost Allocation", "แบ่งค่าไฟตามผู้เช่า/แผนก/โซน อัตโนมัติ", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("5.2.2 Demand Response & Peak Management", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FA-DR01", "Peak Demand Forecasting", "คาดการณ์ peak demand 24–72 ชม. (Prophet / LSTM)", "P0"],
        ["FA-DR02", "Automated Peak Shaving", "ลด peak อัตโนมัติ: ลด HVAC, ปิดไฟ, เลื่อน EV charging", "P1"],
        ["FA-DR03", "Demand Charge Optimization", "ลดค่า demand charge โดยกระจาย load ไม่ให้ peak เกินเกณฑ์", "P1"],
        ["FA-DR04", "Utility Rate Comparison", "เปรียบเทียบแผนค่าไฟ แนะนำแผนประหยัดสุด", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("5.2.3 Renewable Energy Integration", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FA-RE01", "Solar PV Monitoring", "real-time generation, daily/monthly yield, performance ratio", "P1"],
        ["FA-RE02", "Self-consumption Optimization", "เลื่อน load ยืดหยุ่น (EV charging, cooling) ไปช่วงแดดดี", "P1"],
        ["FA-RE03", "BESS Management", "charge จาก solar/off-peak, discharge ตอน peak, automated scheduling", "P2"],
        ["FA-RE04", "Net Metering Dashboard", "import/export energy, net billing, REC accumulation", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("5.2.4 Building Intelligence", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FA-BI01", "HVAC Optimization", "แนะนำ setpoint ตาม occupancy, weather, energy price", "P1"],
        ["FA-BI02", "Lighting Automation", "on/off อัตโนมัติจาก occupancy + daylight sensor", "P2"],
        ["FA-BI03", "Indoor Air Quality (IAQ)", "CO2, PM2.5, Temperature, Humidity ควบคู่กับ energy", "P2"],
        ["FA-BI04", "Digital Twin (Visual)", "3D floor plan + heatmap การใช้พลังงาน real-time", "P3"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("5.2.5 Compliance & Sustainability", level=3)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["FA-CS01", "Carbon Accounting", "Scope 1 & 2 GHG emissions ตาม GHG Protocol", "P1"],
        ["FA-CS02", "ESG Report Generator", "energy, carbon intensity, renewable %, waste reduction", "P2"],
        ["FA-CS03", "LEED / TREES Tracker", "ติดตามคะแนน green building ตาม LEED/TREES", "P3"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  6. CROSS-DOMAIN INTELLIGENCE
# ═══════════════════════════════════════════════════════════════
doc.add_heading("6. โมดูล 4 — Cross-Domain Intelligence", level=1)
add_para("ฟังก์ชันที่เชื่อมโยง 3 โดเมนเข้าด้วยกัน ซึ่งเป็นจุดแข่งขันหลักของระบบ", bold=True)

add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["CD-01", "Unified Energy Dashboard", "ภาพรวมพลังงานทั้งองค์กร: Building + EV + Fleet + Solar ในหน้าเดียว", "P0"],
        ["CD-02", "Site Power Budget Controller", "Building load สูง → ลด EV charging; Solar surplus → เพิ่ม EV charging อัตโนมัติ", "P0"],
        ["CD-03", "Fleet-Charger Matching", "รถ EV เข้า depot → จับคู่สถานีชาร์จอัตโนมัติ ตาม SoC, departure time, availability", "P1"],
        ["CD-04", "Energy Arbitrage Engine", "ซื้อไฟตอนถูก เก็บใน BESS/EV → ใช้ตอนแพง; coordinate V2G + solar + demand", "P2"],
        ["CD-05", "Scenario Simulator", "What-if: เพิ่มสถานี X จุด / เปลี่ยนฟลีทเป็น EV / ติด solar Y kW → ผลกระทบ", "P2"],
        ["CD-06", "Anomaly Correlation", "ตรวจพบ facility spike + charger fault + fleet delay เกิดจากเหตุเดียวกัน", "P2"],
    ],
    col_widths=[1.5, 3.5, 9.5, 1.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  7. PLATFORM FOUNDATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading("7. โมดูล 5 — Platform Foundation", level=1)

doc.add_heading("7.1 Authentication & Authorization", level=2)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["PF-A01", "Multi-tenant", "รองรับหลายองค์กรบนระบบเดียวกัน แยก data isolation", "P0"],
        ["PF-A02", "RBAC", "บทบาท: Super Admin, Org Admin, Site Manager, Fleet Manager, Operator, Viewer", "P0"],
        ["PF-A03", "SSO Integration", "Google Workspace, Microsoft Entra ID, LDAP", "P1"],
        ["PF-A04", "API Key Management", "ออก API key สำหรับ 3rd party + rate limiting", "P1"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("7.2 Notification & Alerting", level=2)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["PF-N01", "Alert Rules Engine", "metric > threshold for duration → action", "P0"],
        ["PF-N02", "Multi-channel", "LINE OA, Email, SMS, In-app notification, Webhook", "P0"],
        ["PF-N03", "Escalation Policy", "ไม่ acknowledge ภายใน X นาที → escalate หัวหน้า", "P1"],
        ["PF-N04", "Alert Analytics", "ประวัติ, วิเคราะห์ alert ที่เกิดบ่อย, ลด false positive", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("7.3 Reporting", level=2)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["PF-R01", "Scheduled Reports", "รายวัน/รายสัปดาห์/รายเดือน ส่งทาง email/LINE", "P0"],
        ["PF-R02", "Custom Report Builder", "drag-and-drop สร้าง report จาก metric ใดก็ได้", "P1"],
        ["PF-R03", "Export", "PDF, Excel, CSV สำหรับทุก report", "P0"],
        ["PF-R04", "Audit Log", "บันทึกทุก action: who did what, when, from where", "P1"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_heading("7.4 Integration & API", level=2)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["PF-I01", "REST API (OpenAPI 3.1)", "Documented API สำหรับทุก resource", "P0"],
        ["PF-I02", "WebSocket Real-time", "Publish real-time: station status, vehicle position, energy readings", "P0"],
        ["PF-I03", "OCPI 2.2 Roaming Hub", "แชร์สถานีชาร์จข้ามเครือข่าย", "P2"],
        ["PF-I04", "ERP Integration", "SAP / Oracle สำหรับ billing, asset management", "P3"],
        ["PF-I05", "BMS/SCADA Gateway", "เชื่อม BMS ผ่าน BACnet/Modbus TCP", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  8. MOBILE
# ═══════════════════════════════════════════════════════════════
doc.add_heading("8. Mobile Application", level=1)
add_table(
    ["รหัส", "ฟังก์ชัน", "รายละเอียด", "Priority"],
    [
        ["MB-01", "Operator Mobile App", "ดูสถานะ, รับ alert, Remote Start/Stop, ดู fleet location", "P1"],
        ["MB-02", "Driver App", "นำทาง, scan QR ชาร์จรถ, proof of delivery, ดู schedule", "P1"],
        ["MB-03", "EV User App (B2C)", "ค้นหาสถานี, จอง, ชำระเงิน, ดูประวัติ", "P2"],
        ["MB-04", "Push Notifications", "ชาร์จเสร็จ, สถานีว่าง, alert, delivery assignment", "P1"],
        ["MB-05", "Offline Mode", "ดูข้อมูล cached, proof of delivery queue เมื่อไม่มีเน็ต", "P2"],
    ],
    col_widths=[2, 3.5, 9, 1.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  9. AI/ML
# ═══════════════════════════════════════════════════════════════
doc.add_heading("9. AI/ML Capabilities", level=1)
add_table(
    ["รหัส", "ฟังก์ชัน", "ML Model", "Input Data", "Priority"],
    [
        ["AI-01", "Energy Demand Forecast", "Prophet / LSTM", "Historical, weather, calendar, occupancy", "P0"],
        ["AI-02", "EV Charging Demand Forecast", "Gradient Boosting", "Session history, time-of-day, weather", "P1"],
        ["AI-03", "Charger Failure Prediction", "Survival Analysis / RF", "Error logs, uptime, meter, temp", "P1"],
        ["AI-04", "Route Time Prediction", "XGBoost", "Trip times, traffic, time-of-day, weather", "P1"],
        ["AI-05", "Optimal Fleet Size", "Monte Carlo Simulation", "Demand, routes, vehicle specs, charging", "P2"],
        ["AI-06", "Anomaly Detection", "Isolation Forest / AE", "All sensor data across domains", "P1"],
        ["AI-07", "Cargo Space Estimation", "YOLOv8 + Depth Estimation", "Camera/LiDAR images of truck interior, load dimensions", "P1"],
    ],
    col_widths=[1.5, 3.5, 3.5, 5.5, 1.5],
)

add_para("")
doc.add_heading("9.1 รายละเอียด AI-07: Cargo Space Estimation", level=2)
add_para(
    "ระบบประเมินพื้นที่บรรทุกสินค้าคงเหลือในตู้คอนเทนเนอร์/ท้ายรถบรรทุก แบบอัตโนมัติ "
    "โดยใช้ Computer Vision วิเคราะห์จากกล้องที่ติดตั้งภายในตู้สินค้าหรือจากภาพถ่ายของคนขับ"
)

add_table(
    ["หัวข้อ", "รายละเอียด"],
    [
        ["วัตถุประสงค์", "ลดเที่ยวรถเปล่า/ไม่เต็มคัน, เพิ่ม load factor เฉลี่ยจาก ~65% เป็น >85%"],
        ["Input Data", "1) ภาพจากกล้อง RGB ติดเพดานตู้สินค้า (หรือภาพถ่ายจาก Driver App)\n"
                       "2) Depth map จาก stereo camera หรือ LiDAR (ถ้ามี)\n"
                       "3) ข้อมูล master: ขนาดตู้คอนเทนเนอร์แต่ละคัน (LxWxH)"],
        ["ML Pipeline", "1) Object Detection (YOLOv8) — ตรวจจับกล่อง/พัสดุแต่ละชิ้นในตู้สินค้า\n"
                        "2) Monocular Depth Estimation (MiDaS / ZoeDepth) — ประมาณความลึกจากภาพ RGB กรณีไม่มี LiDAR\n"
                        "3) 3D Volume Reconstruction — สร้าง point cloud ของพื้นที่ที่ถูกครอบครอง\n"
                        "4) Free Space Calculation — คำนวณปริมาตรว่าง (m³) และ % ที่เหลือ"],
        ["Output", "• Volume utilization (%) — สัดส่วนพื้นที่ที่ใช้ไปแล้ว\n"
                   "• Remaining capacity (m³) — ปริมาตรว่างคงเหลือ\n"
                   "• Load level (0–5) — อัปเดตค่า loadLevel อัตโนมัติ แทนการกรอกมือ\n"
                   "• Packing suggestion — แนะนำตำแหน่งวางสินค้าถัดไป"],
        ["Integration กับ Fleet", "• FL-D01 Smart Dispatch: เลือกรถที่มี free space มากพอสำหรับ order ใหม่\n"
                                  "• FL-D02 Load Consolidation: รวม shipments เข้ารถที่ยังมีที่ว่าง\n"
                                  "• FL-B04 ข้อมูลสินค้า: อัปเดต loadPercent แบบ real-time จาก vision"],
        ["Accuracy Target", "Volume estimation error < ±10% เทียบกับ manual measurement"],
        ["Hardware Options", "Tier 1: กล้อง RGB + MiDaS (ต้นทุนต่ำ, ±15%)\n"
                            "Tier 2: Stereo camera เช่น Intel RealSense (±10%)\n"
                            "Tier 3: LiDAR เช่น Livox Mid-360 (±5%, ต้นทุนสูง)"],
        ["Edge Deployment", "Inference บน edge device (NVIDIA Jetson Orin Nano) ติดที่รถ, "
                           "ส่งเฉพาะผลลัพธ์ (JSON) กลับ server ผ่าน MQTT เพื่อลด bandwidth"],
    ],
    col_widths=[3.5, 13],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  10. NON-FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("10. Non-Functional Requirements", level=1)

doc.add_heading("10.1 Performance", level=2)
add_table(
    ["ข้อกำหนด", "เกณฑ์"],
    [
        ["Dashboard Load Time", "< 2 วินาที (First Contentful Paint)"],
        ["Real-time Data Latency", "< 3 วินาที (sensor → dashboard)"],
        ["API Response Time (p95)", "< 200ms (read), < 500ms (write)"],
        ["Map Rendering", "500 markers @ 60 FPS"],
        ["Concurrent Users", "500 users พร้อมกัน"],
    ],
    col_widths=[5, 12],
)

doc.add_heading("10.2 Availability & Reliability", level=2)
add_table(
    ["ข้อกำหนด", "เกณฑ์"],
    [
        ["Uptime SLA", "99.5% (ยกเว้น planned maintenance)"],
        ["Data Retention", "Raw data 12 เดือน, Aggregated 5 ปี"],
        ["Backup", "Daily backup, RPO < 1 ชม., RTO < 4 ชม."],
        ["Disaster Recovery", "Secondary region failover (optional)"],
    ],
    col_widths=[5, 12],
)

doc.add_heading("10.3 Security", level=2)
add_table(
    ["ข้อกำหนด", "เกณฑ์"],
    [
        ["Authentication", "JWT + Refresh Token, MFA support"],
        ["Data Encryption", "TLS 1.3 in-transit, AES-256 at-rest"],
        ["OWASP Compliance", "ผ่าน OWASP Top 10 assessment"],
        ["Penetration Testing", "ผ่านทดสอบจากทีมภายนอกก่อน go-live"],
        ["PDPA Compliance", "จัดเก็บตาม พ.ร.บ. คุ้มครองข้อมูลส่วนบุคคล"],
    ],
    col_widths=[5, 12],
)

doc.add_heading("10.4 Scalability", level=2)
add_table(
    ["ข้อกำหนด", "เกณฑ์"],
    [
        ["Horizontal Scaling", "เพิ่ม instance ไม่ต้อง downtime (K8s HPA)"],
        ["Database Scaling", "TimescaleDB hypertable partitioning"],
        ["Message Queue", "MQTT cluster > 10,000 concurrent devices"],
    ],
    col_widths=[5, 12],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  11. ROADMAP
# ═══════════════════════════════════════════════════════════════
doc.add_heading("11. แผนการพัฒนา (Implementation Roadmap)", level=1)

phases = [
    ("Phase 1 — Foundation (เดือนที่ 1–3)", [
        "ระบบ Authentication, RBAC, Multi-tenant",
        "EV Charging: OCPP 2.0.1 core, Remote Start/Stop, Session Lifecycle",
        "Fleet: Real-time tracking, Fleet Utilization Dashboard",
        "Facilities: Sub-metering integration, Real-time monitoring",
        "Unified Dashboard (CD-01)",
        "Alert engine + LINE/Email notification",
    ]),
    ("Phase 2 — Intelligence (เดือนที่ 4–6)", [
        "Smart Charging: Load Balancing, TOU Optimization, Charging Profiles",
        "Fleet: Route Optimization, Smart Dispatch, Charge Planning (EV fleet)",
        "Facilities: Peak Demand Forecasting, Demand Charge Optimization",
        "Cross-Domain: Site Power Budget Controller, Fleet-Charger Matching",
        "AI: Energy Demand Forecast, Anomaly Detection",
        "Mobile App v1 (Operator + Driver)",
    ]),
    ("Phase 3 — Advanced (เดือนที่ 7–9)", [
        "Dynamic Pricing, Payment Integration, CDR (OCPI)",
        "V2G Readiness, Battery Energy Storage management",
        "Solar PV monitoring, Self-consumption optimization",
        "Predictive Maintenance, Health Score",
        "Carbon Accounting, ESG Report",
        "Scenario Simulator (CD-05)",
    ]),
    ("Phase 4 — Scale & Polish (เดือนที่ 10–12)", [
        "OCPI 2.2 Roaming Hub",
        "ERP/BMS Integration",
        "Digital Twin (Visual)",
        "EV User App (B2C)",
        "Load testing, Penetration testing, PDPA audit",
        "Documentation, training, go-live",
    ]),
]

for phase_title, items in phases:
    doc.add_heading(phase_title, level=2)
    for item in items:
        add_para(f"• {item}")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  12. JUDGING CRITERIA
# ═══════════════════════════════════════════════════════════════
doc.add_heading("12. เกณฑ์การตัดสิน (สำหรับการแข่งขัน)", level=1)
add_table(
    ["เกณฑ์", "ฟังก์ชันที่ตอบโจทย์", "น้ำหนัก"],
    [
        ["Innovation", "Cross-Domain Intelligence, V2G, Energy Arbitrage, Scenario Simulator", "25%"],
        ["Technical Completeness", "OCPP 2.0.1, OCPI 2.2, ISO 15118, ML pipeline", "25%"],
        ["Practical Impact", "ลดค่าไฟ (TOU + Peak Shaving + Solar), ลดต้นทุน fleet, ลด carbon", "20%"],
        ["User Experience", "Unified Dashboard, Mobile App, Real-time Map, Kiosk Mode", "15%"],
        ["Scalability & Architecture", "Multi-tenant, Kubernetes, TimescaleDB, MQTT cluster", "15%"],
    ],
    col_widths=[3.5, 8, 2],
)

# ═══════════════════════════════════════════════════════════════
#  13. SUMMARY TABLE
# ═══════════════════════════════════════════════════════════════
doc.add_heading("13. ตารางสรุปฟังก์ชันทั้งหมด", level=1)
add_table(
    ["โดเมน", "Baseline (มีแล้ว)", "Enhanced (ใหม่)", "รวม"],
    [
        ["EV Charging", "6", "19", "25"],
        ["Fleet Optimization", "5", "18", "23"],
        ["Facilities Energy", "5", "18", "23"],
        ["Cross-Domain", "0", "6", "6"],
        ["Platform Foundation", "0", "15", "15"],
        ["Mobile", "0", "5", "5"],
        ["AI/ML", "0", "7", "7"],
        ["รวมทั้งหมด", "16", "88", "104"],
    ],
    col_widths=[5, 3, 3, 3],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  14. GLOSSARY
# ═══════════════════════════════════════════════════════════════
doc.add_heading("14. คำจำกัดความและคำย่อ", level=1)
add_table(
    ["คำย่อ", "ความหมาย"],
    [
        ["OCPP", "Open Charge Point Protocol — มาตรฐานสื่อสารระหว่าง CSMS กับสถานีชาร์จ"],
        ["OCPI", "Open Charge Point Interface — มาตรฐาน roaming ระหว่างเครือข่ายชาร์จ"],
        ["ISO 15118", "มาตรฐาน Plug & Charge (ยืนยันตัวตนผ่านสาย EV)"],
        ["CSMS", "Charging Station Management System"],
        ["SoC", "State of Charge — ระดับแบตเตอรี่ (%)"],
        ["SoH", "State of Health — สุขภาพแบตเตอรี่ (%)"],
        ["TOU", "Time of Use — อัตราค่าไฟตามช่วงเวลา"],
        ["V2G", "Vehicle-to-Grid — เทคโนโลยีจ่ายไฟจาก EV คืนกริด"],
        ["BESS", "Battery Energy Storage System"],
        ["PUE", "Power Usage Effectiveness — ดัชนีประสิทธิภาพพลังงานอาคาร"],
        ["MTBF", "Mean Time Between Failures"],
        ["CDR", "Charge Detail Record — บันทึกธุรกรรมการชาร์จ"],
        ["GHG", "Greenhouse Gas"],
        ["PDPA", "พ.ร.บ. คุ้มครองข้อมูลส่วนบุคคล พ.ศ. 2562"],
        ["HOS", "Hours of Service — กฎเกณฑ์ชั่วโมงทำงานคนขับ"],
        ["TCO", "Total Cost of Ownership"],
        ["HPA", "Horizontal Pod Autoscaler (Kubernetes)"],
        ["RPO/RTO", "Recovery Point Objective / Recovery Time Objective"],
    ],
    col_widths=[3, 14],
)

# ── Footer note ──
doc.add_paragraph()
add_para(
    "เอกสารฉบับนี้จัดทำเพื่อใช้เป็นสเปคระบบสำหรับการแข่งขัน โดยอ้างอิงจากระบบ SIAM EV CSMS "
    "ที่พัฒนาแล้วบางส่วน และเพิ่มเติมฟังก์ชันที่สามารถพัฒนาได้จริงด้วยเทคโนโลยีที่มีอยู่ในปัจจุบัน",
    size=12,
)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), "TOR_SIAM_EV_PLATFORM.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
